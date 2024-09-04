import os
from flask import Flask, request, render_template, send_file
from docx import Document
from docx.shared import Pt
from huggingface_hub import InferenceClient

# Initialize the InferenceClient
client = InferenceClient(token="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")

app = Flask(__name__)

def create_code_style(doc):
    styles = doc.styles
    new_style = styles.add_style('CodeStyle', 1)
    new_style.font.name = 'Courier New'
    new_style.font.size = Pt(10)
    return new_style

def add_code_paragraph(doc, text):
    paragraph = doc.add_paragraph(style='CodeStyle')
    run = paragraph.add_run(f"```\n{text}\n```")
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

def create_project_docx(domain_name):
    doc = Document()
    create_code_style(doc)

    doc.add_heading(f'Internship Project Tasks for {domain_name}', 0)

    # Use Hugging Face Inference API to generate tasks
    task_1_response = client.chat_completion(
        messages=[{"role": "user", "content": f"Create a basic application in {domain_name}"}],
        max_tokens=100
    )
    task_1_content = task_1_response['choices'][0]['message']['content']

    task_2_response = client.chat_completion(
        messages=[{"role": "user", "content": f"Develop a simple project that demonstrates {domain_name} principles"}],
        max_tokens=100
    )
    task_2_content = task_2_response['choices'][0]['message']['content']

    doc.add_heading('Task 1', level=1)
    doc.add_paragraph(task_1_content)
    doc.add_paragraph("Sample Code:")
    add_code_paragraph(doc, "# Sample code for Task 1\nprint('Hello, World!')")

    doc.add_heading('Task 2', level=1)
    doc.add_paragraph(task_2_content)
    doc.add_paragraph("Sample Code:")
    add_code_paragraph(doc, "x = 10\ny = 20\nprint(x + y)")

    # Save the document in a temporary directory
    temp_dir = os.path.join(os.getcwd(), 'temp')
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    file_name = os.path.join(temp_dir, f'{domain_name}_project_tasks.docx')
    doc.save(file_name)

    return file_name

@app.route('/', methods=['GET', 'POST'])
def chat():
    if request.method == 'POST':
        domain_name = request.form['domain_name']
        file_name = create_project_docx(domain_name)
        return send_file(file_name, as_attachment=True)

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
